"""Document text extraction utility.

Supports PDF, DOCX and plain text sources for the recruiter workflow so the
pipeline can be fed with whatever format the recruiter uploads.
"""

from io import BytesIO
from typing import Union

from fastapi import UploadFile, HTTPException

from api.utils.logging import logger


SUPPORTED_TEXT_MIMES = {
    "text/plain",
    "text/markdown",
    "text/csv",
    "application/json",
}

PDF_MIMES = {"application/pdf"}

DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}


def _pdf_bytes_to_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as e:  # pragma: no cover
        raise HTTPException(
            status_code=500, detail=f"PDF parser not installed: {e}"
        )

    try:
        reader = PdfReader(BytesIO(data))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception as e:
                logger.warning(f"PDF page extraction failed: {e}")
        return "\n".join(p for p in parts if p).strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read PDF: {e}")


def _docx_bytes_to_text(data: bytes) -> str:
    try:
        import docx  # python-docx
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=(
                "DOCX parser not installed. Add 'python-docx' to backend "
                f"dependencies and reinstall. ({e})"
            ),
        )

    try:
        document = docx.Document(BytesIO(data))
        lines = [p.text for p in document.paragraphs if p.text and p.text.strip()]
        # Also include table cell text (common in resumes/JDs)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        lines.append(cell.text.strip())
        return "\n".join(lines).strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read DOCX: {e}")


async def extract_text_from_upload(upload: UploadFile) -> str:
    """Extract plain text from an UploadFile (PDF / DOCX / text)."""
    if upload is None:
        return ""

    filename = (upload.filename or "").lower()
    content_type = (upload.content_type or "").lower()
    data = await upload.read()

    if not data:
        return ""

    if content_type in PDF_MIMES or filename.endswith(".pdf"):
        return _pdf_bytes_to_text(data)

    if content_type in DOCX_MIMES or filename.endswith(".docx") or filename.endswith(".doc"):
        return _docx_bytes_to_text(data)

    if content_type in SUPPORTED_TEXT_MIMES or content_type.startswith("text/") or filename.endswith((".txt", ".md")):
        try:
            return data.decode("utf-8", errors="ignore").strip()
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to decode text: {e}")

    # Last resort: try utf-8 decode (some uploads have empty content_type)
    try:
        return data.decode("utf-8", errors="ignore").strip()
    except Exception:
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type: {content_type or filename}",
        )


async def extract_text(source: Union[UploadFile, str, None]) -> str:
    """Normalize any input (UploadFile or raw string) into plain text."""
    if source is None:
        return ""
    if isinstance(source, str):
        return source.strip()
    return await extract_text_from_upload(source)
